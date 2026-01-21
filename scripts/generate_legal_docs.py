from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    return heading
def add_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p
def add_bullet_list(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')
    return
def create_privacy_policy_en():
    doc = Document()
    title = doc.add_heading('Privacy Policy', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('NIGINart Social Analytics Dashboard')
    doc.add_paragraph('Last Updated: January 9, 2025')
    doc.add_paragraph()
    doc.add_paragraph(
        'This Privacy Policy describes how NIGINart Social Analytics Dashboard '
        '("we", "our", or "the Service") collects, uses, and protects information '
        'when you use our social media analytics application.'
    )
    add_heading(doc, '1. Information We Collect', 2)
    add_heading(doc, '1.1 Information You Provide', 3)
    add_bullet_list(doc, [
        'Account Information: Social media account identifiers (usernames, channel IDs) '
        'that you choose to connect to our Service.',
        'OAuth Tokens: Access tokens obtained through OAuth 2.0 authorization with platforms '
        'like TikTok, Instagram, and others. These tokens are encrypted and stored securely.'
    ])
    add_heading(doc, '1.2 Information Collected Automatically', 3)
    add_bullet_list(doc, [
        'Social Media Metrics: Public metrics from your connected accounts including '
        'followers count, engagement rates, views, likes, comments, and shares.',
        'Video/Post Metadata: Titles, descriptions, publication dates, and performance metrics.',
        'Historical Data: We store historical metrics to provide trend analysis.'
    ])
    add_heading(doc, '1.3 Information We Do NOT Collect', 3)
    add_bullet_list(doc, [
        'Private messages or direct communications',
        'Personal contact information (email, phone) unless explicitly provided',
        'Financial or payment information',
        'Data from accounts you haven\'t explicitly connected'
    ])
    add_heading(doc, '2. How We Use Your Information', 2)
    add_bullet_list(doc, [
        'Analytics Display: To show you metrics, charts, and insights about your social media performance.',
        'Trend Analysis: To calculate growth rates, engagement trends, and performance comparisons.',
        'Data Export: To enable you to export your analytics data in various formats.',
        'Service Improvement: To improve the functionality and user experience of our Service.'
    ])
    add_heading(doc, '3. Data Storage and Security', 2)
    add_bullet_list(doc, [
        'Encryption: All OAuth access tokens are encrypted using industry-standard Fernet symmetric encryption.',
        'Database Security: Data is stored in PostgreSQL with access controls and regular backups.',
        'Access Control: Only authorized personnel have access to the database.',
        'No Plain Text Storage: Sensitive credentials are never stored in plain text.'
    ])
    add_heading(doc, '4. Data Sharing and Disclosure', 2)
    doc.add_paragraph('We do NOT sell, rent, or share your personal data with third parties except:')
    add_bullet_list(doc, [
        'With Your Consent: When you explicitly authorize sharing.',
        'Legal Requirements: When required by law, court order, or governmental authority.',
        'Service Providers: With hosting providers necessary to operate the Service.'
    ])
    add_heading(doc, '5. Third-Party Platform Data', 2)
    doc.add_paragraph('Our Service integrates with the following platforms through their official APIs:')
    add_bullet_list(doc, [
        'TikTok (via TikTok for Developers API)',
        'Instagram (via Facebook Graph API)',
        'YouTube (via YouTube Data API)',
        'VK (via VK API)',
        'Telegram (via Telegram API)'
    ])
    add_heading(doc, '6. Data Retention', 2)
    add_bullet_list(doc, [
        'Metrics Data: Retained for the lifetime of your account for trend analysis.',
        'OAuth Tokens: Retained until you disconnect your account or they expire.',
        'Account Deletion: Upon request, we will delete all data within 30 days.'
    ])
    add_heading(doc, '7. Your Rights', 2)
    add_bullet_list(doc, [
        'Access: Request a copy of the data we store about your accounts.',
        'Correction: Request correction of inaccurate data.',
        'Deletion: Request deletion of your data and connected accounts.',
        'Revocation: Disconnect any social media account at any time.',
        'Export: Export your analytics data in standard formats (CSV, JSON, Excel).'
    ])
    add_heading(doc, '8. Children\'s Privacy', 2)
    doc.add_paragraph(
        'Our Service is not intended for users under 18 years of age. '
        'We do not knowingly collect data from minors.'
    )
    add_heading(doc, '9. International Data Transfers', 2)
    doc.add_paragraph(
        'Your data may be processed on servers located in different jurisdictions. '
        'We ensure appropriate safeguards are in place.'
    )
    add_heading(doc, '10. Changes to This Policy', 2)
    doc.add_paragraph(
        'We may update this Privacy Policy from time to time. '
        'We will notify you by updating the "Last Updated" date.'
    )
    add_heading(doc, '11. Contact Us', 2)
    doc.add_paragraph('Email: privacy@niginart.com')
    doc.add_paragraph()
    doc.add_paragraph('© 2025 NIGINart. All rights reserved.')
    return doc
def create_privacy_policy_ru():
    doc = Document()
    title = doc.add_heading('Политика конфиденциальности', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('NIGINart Social Analytics Dashboard')
    doc.add_paragraph('Последнее обновление: 9 января 2025 г.')
    doc.add_paragraph()
    doc.add_paragraph(
        'Настоящая Политика конфиденциальности описывает, как NIGINart Social Analytics Dashboard '
        '(«мы», «наш» или «Сервис») собирает, использует и защищает информацию '
        'при использовании нашего приложения для аналитики социальных сетей.'
    )
    add_heading(doc, '1. Информация, которую мы собираем', 2)
    add_heading(doc, '1.1 Информация, которую вы предоставляете', 3)
    add_bullet_list(doc, [
        'Информация об аккаунтах: Идентификаторы аккаунтов социальных сетей (имена пользователей, ID каналов), '
        'которые вы решили подключить к нашему Сервису.',
        'OAuth-токены: Токены доступа, полученные через авторизацию OAuth 2.0 с платформами '
        'TikTok, Instagram и другими. Эти токены шифруются и хранятся безопасно.'
    ])
    add_heading(doc, '1.2 Информация, собираемая автоматически', 3)
    add_bullet_list(doc, [
        'Метрики социальных сетей: Публичные метрики ваших подключённых аккаунтов, включая '
        'количество подписчиков, уровень вовлечённости, просмотры, лайки, комментарии и репосты.',
        'Метаданные видео/постов: Названия, описания, даты публикации и показатели эффективности.',
        'Исторические данные: Мы храним исторические метрики для анализа трендов.'
    ])
    add_heading(doc, '1.3 Информация, которую мы НЕ собираем', 3)
    add_bullet_list(doc, [
        'Личные сообщения или прямые переписки',
        'Персональные контактные данные (email, телефон), если они не предоставлены явно',
        'Финансовая или платёжная информация',
        'Данные аккаунтов, которые вы явно не подключали'
    ])
    add_heading(doc, '2. Как мы используем вашу информацию', 2)
    add_bullet_list(doc, [
        'Отображение аналитики: Для показа метрик, графиков и инсайтов о вашей эффективности в социальных сетях.',
        'Анализ трендов: Для расчёта темпов роста, трендов вовлечённости и сравнения показателей.',
        'Экспорт данных: Для возможности экспорта аналитических данных в различных форматах.',
        'Улучшение сервиса: Для улучшения функциональности и пользовательского опыта.'
    ])
    add_heading(doc, '3. Хранение и безопасность данных', 2)
    add_bullet_list(doc, [
        'Шифрование: Все OAuth-токены шифруются с использованием стандартного симметричного шифрования Fernet.',
        'Безопасность базы данных: Данные хранятся в PostgreSQL с контролем доступа и регулярным резервным копированием.',
        'Контроль доступа: Только авторизованный персонал имеет доступ к базе данных.',
        'Отсутствие хранения в открытом виде: Конфиденциальные учётные данные никогда не хранятся в открытом виде.'
    ])
    add_heading(doc, '4. Передача и раскрытие данных', 2)
    doc.add_paragraph('Мы НЕ продаём, не сдаём в аренду и не передаём ваши персональные данные третьим лицам, за исключением:')
    add_bullet_list(doc, [
        'С вашего согласия: Когда вы явно разрешаете передачу.',
        'Требования законодательства: По требованию закона, судебного приказа или государственных органов.',
        'Поставщики услуг: Хостинг-провайдерам, необходимым для работы Сервиса.'
    ])
    add_heading(doc, '5. Данные сторонних платформ', 2)
    doc.add_paragraph('Наш Сервис интегрируется со следующими платформами через их официальные API:')
    add_bullet_list(doc, [
        'TikTok (через TikTok for Developers API)',
        'Instagram (через Facebook Graph API)',
        'YouTube (через YouTube Data API)',
        'ВКонтакте (через VK API)',
        'Telegram (через Telegram API)'
    ])
    add_heading(doc, '6. Сроки хранения данных', 2)
    add_bullet_list(doc, [
        'Данные метрик: Хранятся в течение всего срока существования вашего аккаунта для анализа трендов.',
        'OAuth-токены: Хранятся до отключения аккаунта или истечения срока действия.',
        'Удаление аккаунта: По запросу мы удалим все данные в течение 30 дней.'
    ])
    add_heading(doc, '7. Ваши права', 2)
    add_bullet_list(doc, [
        'Доступ: Запросить копию данных, которые мы храним о ваших аккаунтах.',
        'Исправление: Запросить исправление неточных данных.',
        'Удаление: Запросить удаление ваших данных и подключённых аккаунтов.',
        'Отзыв: Отключить любой аккаунт социальной сети в любое время.',
        'Экспорт: Экспортировать аналитические данные в стандартных форматах (CSV, JSON, Excel).'
    ])
    add_heading(doc, '8. Конфиденциальность детей', 2)
    doc.add_paragraph(
        'Наш Сервис не предназначен для пользователей младше 18 лет. '
        'Мы сознательно не собираем данные несовершеннолетних.'
    )
    add_heading(doc, '9. Международная передача данных', 2)
    doc.add_paragraph(
        'Ваши данные могут обрабатываться на серверах в различных юрисдикциях. '
        'Мы обеспечиваем надлежащие меры защиты.'
    )
    add_heading(doc, '10. Изменения в политике', 2)
    doc.add_paragraph(
        'Мы можем периодически обновлять эту Политику конфиденциальности. '
        'Мы уведомим вас, обновив дату «Последнее обновление».'
    )
    add_heading(doc, '11. Контакты', 2)
    doc.add_paragraph('Email: privacy@niginart.com')
    doc.add_paragraph()
    doc.add_paragraph('© 2025 NIGINart. Все права защищены.')
    return doc
def create_terms_of_use_en():
    doc = Document()
    title = doc.add_heading('Terms of Use', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('NIGINart Social Analytics Dashboard')
    doc.add_paragraph('Last Updated: January 9, 2025')
    doc.add_paragraph()
    doc.add_paragraph(
        'Welcome to NIGINart Social Analytics Dashboard ("Service"). '
        'By accessing or using our Service, you agree to be bound by these Terms of Use ("Terms"). '
        'Please read them carefully.'
    )
    add_heading(doc, '1. Description of Service', 2)
    doc.add_paragraph('NIGINart Social Analytics Dashboard is a social media analytics platform that:')
    add_bullet_list(doc, [
        'Collects and aggregates metrics from your connected social media accounts',
        'Provides visual dashboards and analytics reports',
        'Tracks historical performance and trends',
        'Enables data export in various formats'
    ])
    add_heading(doc, '2. Eligibility', 2)
    doc.add_paragraph('To use our Service, you must:')
    add_bullet_list(doc, [
        'Be at least 18 years of age',
        'Have the legal authority to bind yourself to these Terms',
        'Own or have authorization to connect the social media accounts you add',
        'Comply with all applicable laws and regulations'
    ])
    add_heading(doc, '3. Account Authorization', 2)
    add_heading(doc, '3.1 OAuth Authorization', 3)
    add_bullet_list(doc, [
        'You authorize us to access specific data as described in the OAuth consent screen',
        'You can revoke access at any time through the Service or directly on the platform',
        'We will only request permissions necessary for analytics functionality'
    ])
    add_heading(doc, '3.2 Your Responsibilities', 3)
    add_bullet_list(doc, [
        'You are responsible for maintaining the security of your connected accounts',
        'You must notify us immediately of any unauthorized use',
        'You must not share access to the Service with unauthorized parties'
    ])
    add_heading(doc, '4. Acceptable Use', 2)
    doc.add_paragraph('You agree NOT to:')
    add_bullet_list(doc, [
        'Use the Service for any unlawful purpose',
        'Attempt to gain unauthorized access to our systems or other users\' data',
        'Use automated scripts or bots to access the Service beyond normal usage',
        'Reverse engineer, decompile, or attempt to extract source code',
        'Resell, redistribute, or commercially exploit the Service without authorization',
        'Use the Service to collect data for spam, harassment, or malicious purposes',
        'Violate any third-party platform\'s Terms of Service through our Service'
    ])
    add_heading(doc, '5. Third-Party Platform Compliance', 2)
    doc.add_paragraph(
        'IMPORTANT: Your use of our Service must comply with the Terms of Service '
        'of each connected platform (TikTok, Instagram, YouTube, VK, Telegram, etc.).'
    )
    add_bullet_list(doc, [
        'Each platform has its own Terms of Service that govern your account',
        'We are not responsible for changes to third-party APIs or Terms',
        'Platform API limitations may affect Service functionality',
        'Violation of platform Terms may result in revocation of API access'
    ])
    add_heading(doc, '6. Data Usage and Privacy', 2)
    doc.add_paragraph('Please review our Privacy Policy to understand how we collect, use, and protect your data.')
    add_bullet_list(doc, [
        'We collect only data necessary for analytics functionality',
        'OAuth tokens are encrypted before storage',
        'We do not sell your data to third parties',
        'You can request data deletion at any time'
    ])
    add_heading(doc, '7. Intellectual Property', 2)
    add_heading(doc, '7.1 Our Property', 3)
    doc.add_paragraph(
        'The Service, including its design, code, features, and documentation, '
        'is owned by NIGINart and protected by intellectual property laws.'
    )
    add_heading(doc, '7.2 Your Content', 3)
    doc.add_paragraph(
        'You retain ownership of all content and data from your social media accounts. '
        'By using the Service, you grant us a limited license to access and process this data '
        'solely to provide analytics services.'
    )
    add_heading(doc, '8. Disclaimer of Warranties', 2)
    doc.add_paragraph(
        'THE SERVICE IS PROVIDED "AS IS" AND "AS AVAILABLE" WITHOUT WARRANTIES OF ANY KIND, '
        'EITHER EXPRESS OR IMPLIED, INCLUDING BUT NOT LIMITED TO:'
    )
    add_bullet_list(doc, [
        'Accuracy or completeness of analytics data',
        'Uninterrupted or error-free operation',
        'Fitness for a particular purpose',
        'Compatibility with third-party platforms'
    ])
    add_heading(doc, '9. Limitation of Liability', 2)
    doc.add_paragraph('To the maximum extent permitted by law:')
    add_bullet_list(doc, [
        'We are not liable for any indirect, incidental, special, or consequential damages',
        'We are not liable for loss of data, profits, or business opportunities',
        'Our total liability shall not exceed the amount paid by you in the past 12 months',
        'We are not responsible for third-party platform outages or data accuracy'
    ])
    add_heading(doc, '10. Indemnification', 2)
    doc.add_paragraph('You agree to indemnify and hold harmless NIGINart from any claims arising from:')
    add_bullet_list(doc, [
        'Your violation of these Terms',
        'Your violation of any third-party rights',
        'Your misuse of the Service',
        'Content from your connected social media accounts'
    ])
    add_heading(doc, '11. Service Modifications and Termination', 2)
    doc.add_paragraph('We reserve the right to:')
    add_bullet_list(doc, [
        'Modify, suspend, or discontinue the Service at any time',
        'Change features or functionality with or without notice',
        'Terminate your access for violation of these Terms',
        'Update these Terms with notice posted on the Service'
    ])
    add_heading(doc, '12. Governing Law', 2)
    doc.add_paragraph(
        'These Terms shall be governed by and construed in accordance with applicable laws. '
        'Any disputes shall be resolved through negotiation or competent courts.'
    )
    add_heading(doc, '13. Contact Information', 2)
    doc.add_paragraph('For questions about these Terms of Use:')
    doc.add_paragraph('Email: legal@niginart.com')
    doc.add_paragraph()
    doc.add_paragraph('© 2025 NIGINart. All rights reserved.')
    return doc
def create_terms_of_use_ru():
    doc = Document()
    title = doc.add_heading('Условия использования', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('NIGINart Social Analytics Dashboard')
    doc.add_paragraph('Последнее обновление: 9 января 2025 г.')
    doc.add_paragraph()
    doc.add_paragraph(
        'Добро пожаловать в NIGINart Social Analytics Dashboard («Сервис»). '
        'Получая доступ к нашему Сервису или используя его, вы соглашаетесь соблюдать настоящие Условия использования («Условия»). '
        'Пожалуйста, внимательно ознакомьтесь с ними.'
    )
    add_heading(doc, '1. Описание Сервиса', 2)
    doc.add_paragraph('NIGINart Social Analytics Dashboard — это платформа аналитики социальных сетей, которая:')
    add_bullet_list(doc, [
        'Собирает и агрегирует метрики ваших подключённых аккаунтов социальных сетей',
        'Предоставляет визуальные дашборды и аналитические отчёты',
        'Отслеживает историческую эффективность и тренды',
        'Позволяет экспортировать данные в различных форматах'
    ])
    add_heading(doc, '2. Требования к пользователям', 2)
    doc.add_paragraph('Для использования нашего Сервиса вы должны:')
    add_bullet_list(doc, [
        'Быть не младше 18 лет',
        'Иметь законные полномочия принять настоящие Условия',
        'Владеть или иметь разрешение на подключение аккаунтов социальных сетей',
        'Соблюдать все применимые законы и правила'
    ])
    add_heading(doc, '3. Авторизация аккаунтов', 2)
    add_heading(doc, '3.1 OAuth-авторизация', 3)
    add_bullet_list(doc, [
        'Вы разрешаете нам доступ к определённым данным, как описано на экране согласия OAuth',
        'Вы можете отозвать доступ в любое время через Сервис или непосредственно на платформе',
        'Мы запрашиваем только разрешения, необходимые для функций аналитики'
    ])
    add_heading(doc, '3.2 Ваши обязанности', 3)
    add_bullet_list(doc, [
        'Вы несёте ответственность за безопасность подключённых аккаунтов',
        'Вы обязаны немедленно уведомить нас о несанкционированном использовании',
        'Вы не должны предоставлять доступ к Сервису неавторизованным лицам'
    ])
    add_heading(doc, '4. Допустимое использование', 2)
    doc.add_paragraph('Вы соглашаетесь НЕ:')
    add_bullet_list(doc, [
        'Использовать Сервис в незаконных целях',
        'Пытаться получить несанкционированный доступ к нашим системам или данным других пользователей',
        'Использовать автоматизированные скрипты или ботов сверх обычного использования',
        'Осуществлять обратную разработку, декомпиляцию или извлечение исходного кода',
        'Перепродавать, распространять или коммерчески использовать Сервис без разрешения',
        'Использовать Сервис для сбора данных для спама, преследования или злонамеренных целей',
        'Нарушать Условия использования сторонних платформ через наш Сервис'
    ])
    add_heading(doc, '5. Соответствие требованиям сторонних платформ', 2)
    doc.add_paragraph(
        'ВАЖНО: Использование нашего Сервиса должно соответствовать Условиям использования '
        'каждой подключённой платформы (TikTok, Instagram, YouTube, VK, Telegram и др.).'
    )
    add_bullet_list(doc, [
        'Каждая платформа имеет собственные Условия использования',
        'Мы не несём ответственности за изменения в API или Условиях сторонних платформ',
        'Ограничения API платформ могут влиять на функциональность Сервиса',
        'Нарушение Условий платформ может привести к отзыву доступа к API'
    ])
    add_heading(doc, '6. Использование данных и конфиденциальность', 2)
    doc.add_paragraph('Ознакомьтесь с нашей Политикой конфиденциальности, чтобы понять, как мы собираем и защищаем ваши данные.')
    add_bullet_list(doc, [
        'Мы собираем только данные, необходимые для функций аналитики',
        'OAuth-токены шифруются перед сохранением',
        'Мы не продаём ваши данные третьим лицам',
        'Вы можете запросить удаление данных в любое время'
    ])
    add_heading(doc, '7. Интеллектуальная собственность', 2)
    add_heading(doc, '7.1 Наша собственность', 3)
    doc.add_paragraph(
        'Сервис, включая его дизайн, код, функции и документацию, '
        'принадлежит NIGINart и защищён законами об интеллектуальной собственности.'
    )
    add_heading(doc, '7.2 Ваш контент', 3)
    doc.add_paragraph(
        'Вы сохраняете право собственности на весь контент и данные из ваших аккаунтов социальных сетей. '
        'Используя Сервис, вы предоставляете нам ограниченную лицензию на доступ и обработку этих данных '
        'исключительно для предоставления услуг аналитики.'
    )
    add_heading(doc, '8. Отказ от гарантий', 2)
    doc.add_paragraph(
        'СЕРВИС ПРЕДОСТАВЛЯЕТСЯ «КАК ЕСТЬ» И «ПО МЕРЕ ДОСТУПНОСТИ» БЕЗ КАКИХ-ЛИБО ГАРАНТИЙ, '
        'ЯВНЫХ ИЛИ ПОДРАЗУМЕВАЕМЫХ, ВКЛЮЧАЯ, НО НЕ ОГРАНИЧИВАЯСЬ:'
    )
    add_bullet_list(doc, [
        'Точность или полнота аналитических данных',
        'Бесперебойная или безошибочная работа',
        'Пригодность для определённой цели',
        'Совместимость со сторонними платформами'
    ])
    add_heading(doc, '9. Ограничение ответственности', 2)
    doc.add_paragraph('В максимальной степени, разрешённой законом:')
    add_bullet_list(doc, [
        'Мы не несём ответственности за косвенные, случайные, особые или последующие убытки',
        'Мы не несём ответственности за потерю данных, прибыли или деловых возможностей',
        'Наша общая ответственность не превышает сумму, уплаченную вами за последние 12 месяцев',
        'Мы не несём ответственности за сбои сторонних платформ или точность данных'
    ])
    add_heading(doc, '10. Возмещение убытков', 2)
    doc.add_paragraph('Вы соглашаетесь возместить ущерб NIGINart от любых претензий, возникающих из:')
    add_bullet_list(doc, [
        'Вашего нарушения настоящих Условий',
        'Вашего нарушения прав третьих лиц',
        'Вашего неправомерного использования Сервиса',
        'Контента из ваших подключённых аккаунтов социальных сетей'
    ])
    add_heading(doc, '11. Изменение и прекращение работы Сервиса', 2)
    doc.add_paragraph('Мы оставляем за собой право:')
    add_bullet_list(doc, [
        'Изменять, приостанавливать или прекращать работу Сервиса в любое время',
        'Изменять функции с уведомлением или без него',
        'Прекратить ваш доступ за нарушение настоящих Условий',
        'Обновлять настоящие Условия с уведомлением на Сервисе'
    ])
    add_heading(doc, '12. Применимое право', 2)
    doc.add_paragraph(
        'Настоящие Условия регулируются и толкуются в соответствии с применимым законодательством. '
        'Любые споры разрешаются путём переговоров или в компетентных судах.'
    )
    add_heading(doc, '13. Контактная информация', 2)
    doc.add_paragraph('По вопросам настоящих Условий использования:')
    doc.add_paragraph('Email: legal@niginart.com')
    doc.add_paragraph()
    doc.add_paragraph('© 2025 NIGINart. Все права защищены.')
    return doc
def main():
    output_dir = Path(__file__).parent.parent / 'docs' / 'legal'
    output_dir.mkdir(parents=True, exist_ok=True)
    documents = {
        'privacy-policy-en.docx': create_privacy_policy_en,
        'privacy-policy-ru.docx': create_privacy_policy_ru,
        'terms-of-use-en.docx': create_terms_of_use_en,
        'terms-of-use-ru.docx': create_terms_of_use_ru,
    }
    for filename, create_func in documents.items():
        doc = create_func()
        filepath = output_dir / filename
        doc.save(filepath)
        print(f"[OK] Created: {filepath}")
    print(f"\nAll documents saved to: {output_dir}")
if __name__ == '__main__':
    main()
